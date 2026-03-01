from docx import Document
import json
import re
import datetime


document = Document('CV_YinongChen.docx')
text = "\n".join([p.text for p in document.paragraphs])
paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip() != ""]
#print(paragraphs)

dict_output = {
    "Name": "",
    "Highest_degree": "",
    "Rank": "",
    "Academic_apointment": "",
    "Ft_Pt": "",
    "gov": 0,
    "teaching": 0,
    "institution": 0,
    "Professional_regestration": "",
    "Professional_organizations": "",
    "Professional_development": "",
    "Summer_work": ""
    
}
current_year = datetime.datetime.now().year

"""
dict_matchings based on this
1. Code: P = Professor ASC = Associate Professor AST = Assistant Professor I = Instructor A = Adjunct O = Other
2. Code: T = Tenured  TT = Tenure Track  NTT = Non-Tenure Track
3. FT = Full-Time Faculty or PT = Part-Time Faculty
4. The level of activity (high, medium or low) should reflect an average over the three years prior to the visit.
"""
dict_matchings = {
    "professor": "P",
    "principal lecture": "PL",
    "associate professor": "ASC",
    "assistant professor": "AST",
    "instructor": "I",
    "adjunct": "A",
    "other": "O",
    "tenured": "T",
    "tenure track": "TT",
    "non-tenure track": "NTT",
    "full time": "FT",
    "part time": "PT"
}


def number_years(experience: str) -> int:

    #format must be 4 digit year - 4 digit year xxxx-xxxx
    experience = experience.replace(" ", "")
    experience = experience.replace("–", "-")
    experience = experience.replace("present", str(current_year))
    #finds the start year and end year in academic experience digit and only 4 (\d{4})
    match = re.search(r'(\d{4})-(\d{4})', experience)
    if match:  
        start_year = int(match.group(1))
        end_year = int(match.group(2))
        #print(f"Start Year: {start_year}, End Year: {end_year}")
        return end_year - start_year
    else:
        return 0

#gets name rank and apointment
def user_intro(split_cat: str) -> None:
        
        dict_output["Name"] = split_cat[0].strip()
        dict_output["Rank"] = split_cat[1].strip()
        dict_output["Academic_apointment"] = split_cat[1].strip()

#need to fix later for accuracy
def get_degree(degree: str) -> None:
        dict_output["Highest_degree"] = degree

def teach_inst_exp(i: int,paragraphs: str) -> None:
        for j in range(i+1, len(paragraphs)):
            
            if(paragraphs[j] == 'Non-academic experience'):
                 break
            split_exp = paragraphs[j].split(",")

            years = number_years(paragraphs[j])
            if split_exp[0].strip() == 'Arizona State University':
                dict_output["institution"] = dict_output["institution"] + years
            dict_output["teaching"] = dict_output["teaching"] + years

        #grabs end of academic experience paragraph to determine if full time or part time
        split_cat = paragraphs[i+1].split(",")
        dict_output["Ft_Pt"] = split_cat[-1].strip()

def gov_ind_practice(i: int,paragraphs: str) -> None:
        for j in range(i+1, len(paragraphs)):
            if(paragraphs[j] == 'Certifications or professional registrations'):
                break
            dict_output["gov"] = dict_output["gov"] + number_years(paragraphs[j])
            print(dict_output["gov"])

def curr_memberships(i: int,paragraphs: str) -> None:
        prof_county = 0
        holdy = ""
        for j in range(i+1, len(paragraphs)):
            
            if(paragraphs[j] == 'Honors and awards'):
                break
            prof_county += 1
            paragraphs[j] = paragraphs[j].replace(",", "")
            split_cat = paragraphs[j].split(" ")
            if holdy == "":
                holdy = split_cat[0].strip()
            else:
                holdy = holdy + "," + split_cat[0].strip()
        if prof_county >= 5:
            dict_output["Professional_organizations"] = "H"
        elif prof_county >= 2:
            dict_output["Professional_organizations"] = "M"
        else:
            dict_output["Professional_organizations"] = "L"
        dict_output["Professional_organizations"] = dict_output["Professional_organizations"] + ", " + holdy

def prof_developments(i: int,paragraphs: str) -> None:
        prof_county = 0
        for j in range(i+1, len(paragraphs)):
            prof_county += 1
        print(prof_county)
        if prof_county >= 5:
            dict_output["Professional_development"] = "H"
        elif prof_county >= 2:
            dict_output["Professional_development"] = "M"
        else:
            dict_output["Professional_development"] = "L"





for i in range(len(paragraphs)):
    if(paragraphs[i] == 'Name' and paragraphs[i+1] != 'Education'):
        user_intro(paragraphs[i+1].split(","))
        
    if(paragraphs[i] == 'Education' and paragraphs[i+1] != 'Academic experience'):
        get_degree(paragraphs[i+1])
         

    if(paragraphs[i] == 'Academic experience' and paragraphs[i+1] != 'Non-academic experience '):
        
        #getting years of experience for teaching and institution. teaching and institution based on academic experience, government ind. based on non-academic experience
        teach_inst_exp(i, paragraphs)
        
    #grabs years of experience for government ind practice based on non-academic experience
    if(paragraphs[i] == 'Non-academic experience' and paragraphs[i+1] != 'Certifications or professional registrations'):
        gov_ind_practice(i, paragraphs)
        
    #CISS P stuff
    if(paragraphs[i] == 'Certifications or professional registrations' and paragraphs[i+1] != 'Current membership in professional organizations'):
        dict_output["Professional_regestration"] = paragraphs[i+1]
    
    #figure out level of activity in professional organizations, professional development, and summer work
    if(paragraphs[i] == 'Current membership in professional organizations' and paragraphs[i+1] != 'Honors and awards'):
        curr_memberships(i, paragraphs)
        

    #might need to chage later      
    if(paragraphs[i] == 'Most recent professional development activities'):
        prof_developments(i, paragraphs)
        

    dict_output["Summer_work"] = "N/A" 
        

        #get more later
        
#get rankings
for i in dict_output:
    for key in dict_matchings:
        if str(dict_output[i]).lower() == key:
            dict_output[i] = dict_matchings[key]

print(dict_output)

doc = Document()
doc.add_paragraph(json.dumps(dict_output, indent=4))
table = doc.add_table(rows=1, cols=12)
table.style = 'Table Grid'
table.cell(0, 0).text = "Faculty Name"
table.cell(0, 1).text = "Highest Degree Earned- Field and Year"
table.cell(0, 2).text = "Rank 1"
table.cell(0, 3).text = "Type of Academic Appointment"
table.cell(0, 4).text = "Govt./Ind. Practice"
table.cell(0, 5).text = "Teaching"
table.cell(0, 6).text = "This Institution"
table.cell(0, 7).text = "Professional Registration/ Certification"
table.cell(0, 8).text = "Professional Organizations"
table.cell(0, 9).text = "This Institution"
table.cell(0, 10).text = "Professional Development"
table.cell(0, 11).text = "Consulting/summer work in industry"

#creates tables here
row = table.add_row().cells
nummy = 0
for i in dict_output:
     row[nummy].text = str(dict_output[i])
     nummy += 1
"""
row[0].text = dict_output["Name"]

row[1].text = str(dict_output.get("Highest_degree", ""))
row[2].text = str(dict_output.get("Rank", ""))
row[3].text = str(dict_output.get("Academic_apointment", ""))
row[4].text = str(dict_output.get("gov", ""))
row[5].text = str(dict_output.get("teaching", ""))
row[6].text = str(dict_output.get("institution", ""))
row[7].text = str(dict_output.get("Professional_registration", ""))
row[8].text = str(dict_output.get("Professional_organizations", ""))
row[9].text = str(dict_output.get("This_institution", ""))
row[10].text = str(dict_output.get("Professional_development", ""))
row[11].text = str(dict_output.get("Summer_work", ""))
"""

doc.save('output.docx')
    
   