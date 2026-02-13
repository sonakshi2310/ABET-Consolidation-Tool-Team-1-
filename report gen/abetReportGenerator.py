"""
abetReportGenerator-template_embed.py
------------------------------------
ABET Assessment Report Generator

- Reads Canvas JSON data from input_jsons/*.json
- Writes DOCX reports to generated_pdfs/<json_stem>_ABET_Report.docx
- Uses embedded DOCX template (Base64) written once to generated_pdfs/_ABET_TEMPLATE.docx
- Uses OpenAI ONLY to fill the final "changes needed..." tail when outcome was not met
- Preserves template formatting by replacing text inside existing template paragraphs AND table cells

IMPORTANT:
- Put your template DOCX Base64 into TEMPLATE_DOCX_B64.
"""

import os
import json
import csv
import time
import re
import html
import base64
from glob import glob
from pathlib import Path
from typing import Optional, Dict, Any, Iterable, Tuple

from docx import Document  # python-docx
from dotenv import load_dotenv
from openai import OpenAI


# CONFIG
load_dotenv()

json_input_glob = "input_jsons/*.json"
out_dir = "generated_pdfs"
os.makedirs(out_dir, exist_ok=True)

output_csv = "ABETReportSummary.csv"
template_path = os.path.join(out_dir, "_ABET_TEMPLATE.docx")

OPENAI_MODEL = "gpt-5.2"

prompt_base = (
    "You are an education assessment expert writing ABET assessment report feedback. "
    "Output ONLY the text that should follow the question: "
    "'If outcome was not met, what changes need to be made to ensure that students can meet this outcome in the future?' "
    "Do not add headings, labels, bullets, or extra lines. 1-3 sentences. Concise and actionable."
)

# TEMPLATE BASE64: Template by prof encoded in base64
TEMPLATE_DOCX_B64 = r'''UEsDBBQABgAIAAAAIQDkJIlMfQEAACkGAAATAAgCW0NvbnRlbnRfVHlwZXNdLnhtbCCiBAIooAACAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAC0lMtqwzAQRfeF/oPRtsRKuiilxMmij2UbaArdKvI4EdULaZzH33fsJKYUJy5NsjHIM/feowHNcLw2OllCiMrZjA3SPkvASpcrO8/Yx/Sld8+SiMLmQjsLGdtAZOPR9dVwuvEQE1LbmLEFon/gPMoFGBFT58FSpXDBCKRjmHMv5JeYA7/t9++4dBbBYg8rDzYaPkEhSo3J85p+b0kC6MiSx21jlZUx4b1WUiDV+dLmv1J6u4SUlHVPXCgfb6iB8daEqnI4YKd7o9EElUMyEQFfhaEuvnIh57mTpSFletymhdMVhZLQ6Cs3H5yEGGnmRqdNxQhl9/xtHLKM6Myn0VwhmElwPg5OxmlMKz8IqKCZ4cFZ2NLMIBD9+YfRWHdCRNxoiOcn2Pp2xwMiCS4BsHPuRFjB7P1iFD/MO0EKyp2KmYbzYzTWnRBIqwi239NfRW1zLJI66wdIqy3849r73VWpe/5PL69JJOuT7wfVWswhb8nm9aIffQMAAP//AwBQSwMEFAAGAAgAAAAhAB6RGrfvAAAATgIAAAsACAJfcmVscy8ucmVscyCiBAIooAACAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACsksFqwzAMQO+D/YPRvVHawRijTi9j0NsY2QcIW0lME9vYatf+/TzY2AJd6WFHy9LTk9B6c5xGdeCUXfAallUNir0J1vlew1v7vHgAlYW8pTF41nDiDJvm9mb9yiNJKcqDi1kVis8aBpH4iJjNwBPlKkT25acLaSIpz9RjJLOjnnFV1/eYfjOgmTHV1mpIW3sHqj1FvoYdus4ZfgpmP7GXMy2Qj8Lesl3EVOqTuDKNain1LBpsMC8lnJFirAoa8LzR6nqjv6fFiYUsCaEJiS/7fGZcElr+54rmGT827yFZtF/hbxucXUHzAQAA//8DAFBLAwQUAAYACAAAACEAUMu/e1ALAACqPQAAEQAAAHdvcmQvZG9jdW1lbnQueG1s7Fvrb9s4Ev9+wP0PRBYF9oA01lu2sc1C1qObw+5e0PTuPtMSbesqiQJFxc3+9Tejhy35FcVp0nbvWsCS+PhxOJwZDoeTn37+nCbknoki5tm7C/VKuSAsC3kUZ8t3F//8GLwdX5BC0iyiCc/Yu4sHVlz8fP3Xv/y0nkY8LFOWSQIQWTFd5+G7i5WU+XQ0KsIVS2lxlcah4AVfyKuQpyO+WMQhG625iEaaoirVWy54yIoCxnNpdk+LiwYu/DwMLRJ0DZ0R0BiFKyok+7zFUJ8MYo4mo/E+kHYGEMxQU/eh9CdDWSOkag/IOAsIqNpDMs9DOjA56zwkbR/JPg9J30can4e0J07pvoDznGVQueAipRI+xXKUUvGpzN8CcE5lPI+TWD4ApmK1MDTOPp1BEfTaIKR69GQEe5TyiCV61KLwdxelyKZN/7eb/kj6tO7fPDY9WDJsWBhuMmKfZVLItq8Ywru6u9cYloprI8ES4CPPilWcb6xDei4aVK5akPtTDLhPk7bdOlcHqtox0+bVy7AFHEJ+s3ZpUlN+GlFVBqwmQmx6DCGhP2ZLSQoSvB34LNZ0mKsOND4tgLYHYIVs4GbRYowbjFG41W7EiQeqVYtTrwrixFvGqgNt4C4xHYCofBKEprd04AO7d7CKSEarp8G1azTCvlTSFS02SlMjLgYaghbR6CDWApbwcGPPEJM9jWnmBvAh7axhvnyeor4XvMy3aPHz0G62JnuNztMTsBqF7xqh4nnE3K1oDpY8Dac3y4wLOk+AIlBfAhpIqhXAXxBkfFSv7HNVjvLTvCwSfIlKgibx4hqcwDmPHvCZQ4UxzamgN6BDSmAoge6BM4mlsIVKLLUmE9dRZrARrqfgcEYfoCGIv2F6403RrcBCU7Fs09sUemxBy0TuN7/tNK6ouBX4KHIawoyhEV1IBojAT+ySxLgKml1RgB8fSmQCLSW/GGFHUfefw8eo+RptUKsfee0UBTC0cnw/sJwLSe5YiBsUUafkrkxh/3/AXrLGqBH2WWRMPEsHjvRZZDf/eiw6zo1+TcWNmaGaY6XLjROTOkyaPgkmljbbIU0xdNub2X6PtGa0/urtFHbo7dfs0nuMWuTlP0oJMs4IrbjPoh6H4adLlKGY6mxcQx5EQ5WqhAQWPxesYOKeXVyTAZj9iXYGktc/unc++UH720DKXoYImhFau5xEckLzPHmA89wSRJ0J1IiIFfEywzowDFEZMlLwpKzcKyJXVJKUMUmKnIXxImYRgX5RQdaxXJEQ2sQRE5UzRviC5OU8iUOyYjSRq0tS0AWTD5cEzopkzZIFFQw+CnxP8LlM+JwmlyQEWSgFvhU8jPHJsvtY8AxVCj8RgMFoHMwbWdBQclEMUCgtMDxQqB2bc1ChvpjUVkZjR68VXVesyaxPhqo7uurr6nPIMF1Dnz1D2U8rmHt3WvI0VddPqNS1v9v9yer3w0votK6oA2RHUSaqEVjOa8rOsEU7TK+lOophTXAf69BrmDPH8I3ghS20vP6NSRGHp5ZLVQwvMGujdGRlbrJCijqA9PEhZ4fE50jX6RAj2UhlUdAhu7Hpuoqrmjta++rGw5q41sycaU8nY/gSHiS85tZ7TpNiOkRfbN0ESXtVW/sI8UO208OT0XTFdkx/R5leTfnzO/mQgO89vadwqPo1LuQtELYUNF/VIpKVad0yTu6Ttl0jPlB3E7VlKpaNOh3O9oNh90U+lDS520Bsh21t0xnr8ZGTYsXXBKNjTLIsRIeBFLKM0A6kZQHOR8gF+HmSJIzC51Arc20rTzAhbwj4MHLFhsPXjieaq2NdDijU1QDxAxkztLGJrsH/xe+VxA9XnjcnihX4pnPGMnB95eVzxO2EHLxBfxkHbSS9aES9rwpDpEX39dlkZuIprOteGlYQeP5L7vxPZfUdTXOQqyL+g53erzXPCHy7YZT5BK4esfgti4dsY4FrGJ71qqp3YNe3x95EG7vf/m76gRUAXOys5xFB9QJPN8dI/tfkrW4GM8McuwPIaNT7wIj9mmrEji344oZUe31DWmvfXbX1RkC6iOdVVODgQst50jyas9c86c39I8YW34s4qomA6n9D1RpHhacEZ79LI1TfYMgAG9hap0n0mW5a/Mr5pw3thlO1WsSikB/4GncenD9tvraVLk/KFK+12/q2oGqS8V9mNEMzUH/9q/5qt7Hu9HAy+LqEJ2DUxKogVzWBvWJNsVvudnrKns3bCFT3CKLaY8110YafkNSPlXJ0d6Kw/m1oDRtWV8TtshJJahruGELdsX03wB5fW0vOlvW4EqKELZCYs8T/lokQY1DL/mG0FvuKdycYXi37YIabru44E22IdfwTM/z3Mp0zgb7RiU27YTw+6m6PapLi+YY1cR450L2QJjnmLNDdhqL/1YVVFeXNW1t58+J6BHuBbpvOTsBEV3xbGVtb3p5md+OBfr/sNsZfRG80G4oMy/gqeuPPTPg/JAjzJ9abGUvgRPgaimO5yti1nd2jj68EjmNhGHQIvwNHCwJ0rb9XfmuP6g0+and393gzM/VAPyNS+xpRm3ai9Yz21t4zLE33du53TVODcSb6c2g/YElfMuL01Q5KOwZ33Y9onGJa1f+RyEUdAo1ORk28iQEeew13IhY1iBwMTwlC5/yeYVhUlHMRh+THXcy67cR65lj9C/LDEqpbnqP5wc5ZyPBd1QmqrJJHRuxflXYltOHbn15C23yJNS0wsDmA59bMmLjeeGf3V4Kx6xtq/+J6h72P8XzQxfX3z/PDWn2z2ESacS0yjjkW8pKsMdsiXNFsyYoqzQKTMuaMpDRi+HpSy2xrYk4ahTp8c3qdFaVgdU7Hxq6EFKPcTEJxXGzIiuubkEUpocfPQ43OERv2uzNA1JSJ6YNEDDr99tPDTmeC6YoauH5Xqs5c+j11Op70pU1Jp3J7qT2ADbbrWdrssYhPxYadzJ8OG/o1FRs6N+BfmA0B7BOV7NAC84hwnpfkgZeEoqxxshaxZASkjOG9e5WCFEUgHAWmGy0wD5FnNCFQkmM2UZWpFIP8Y8oPymCcAqESAwJ4K1ICqdv8pKsh+0ZgzhzLx7l9Hxw9rEVg+SRmE/rA6xD4hew5zcABvMGkWVPXduy75ftjZaz3vb5vgjdVPJdlEQNX6JYu2Uww+qkyy/LaY0Uo4jnaN0bYYC6BXNFGspC+buMrQn6Bo1fEz4cEYe21jeLFAsoWgqcAyguGjTnAi15KXrdP0Te/RzZqT9MNOGZ9P2ajkmcwlDeVeg+5mbNdy5o5wZBwwDcyR7dJjyTtnW7ferWmLWc8T9glyRPQ+drwgQDkTMiHK7IV6zXvyOlOiuaK3reWErwGXgqy5iKJKiwcfM1IEt+zK4LWmtFwtQG4xPzPegRoOIrxThr9g6a6ugqv/A/aZGs2w2BO5jYls1vWzdzsVEQjPFBUCZ5sW3NF7sq82jwfkOowoXHaJJcKBqrCMmDJEDOvuGpgucYQ/+GbN/M6qEWGfxIBa/n3MloiL4fZc3c8Gc/sITmL34iKeGVl8OrFB4+Bidp5SMFVAJEEg5sAE8ocU3M4lH5iIK9hjKaxEf11nLTCVFvSoqM0KPsIt+YlaEPVPW45+5+Ws1dko6gJW2K6cc9mV7nHjfHPgdowBm0dJJN2oDmmZe0GhCx3YvneTkDtMZ/2y/H+iOiMHd/1nSHXfY+KzpentQDP+lYcJwI9byydOGqbk5wv7/4gddRZ06q0xOkK3s2x0SYFLH+jCCl5jpHpuomIlyu5/ZxzKXm6/a4Di+3XioFhFJsr6gXn1RSbz2Upq89NGDIBw9ZqPLapiiMe4n1ww4XbWIYrjHW0p8l64tVr/Wc1o+0fWV//FwAA//8DAFBLAwQUAAYACAAAACEA1+YInhsBAABEBAAAHAAIAXdvcmQvX3JlbHMvZG9jdW1lbnQueG1sLnJlbHMgogQBKKAAAQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACsk81OwzAQhO9IvIO1d+KkQEGoTi8IqVcIElc32fyI2I7sDZC3x6Jqkooq6sHHHWtnPnu9m+2PatkXWtcYLSCJYmCoc1M0uhLwnr3cPAJzJHUhW6NRwIAOtun11eYVW0m+ydVN55h30U5ATdQ9ce7yGpV0kelQ+5PSWCXJl7bincw/ZYV8FcdrbucekJ54sl0hwO6KW2DZ0OEl3qYsmxyfTd4r1HQmgjsaWn8BlklbIQk41JH3AX4+/iFkPPlenNL/yoOYLDGsQjLoXu3R+vFOHKO0BJGEhMh7R0Z9+LQRIoomlTeEavFJ1iFpSqMpk/t2NppRWoK4Dwnxjfs3JPJTmP3PmbgEchd0R/5RHJUjAj/Z/fQXAAD//wMAUEsDBBQABgAIAAAAIQCWta3i8QUAAFAbAAAVAAAAd29yZC90aGVtZS90aGVtZTEueG1s7FlLbxNHHL9X6ncY7R38iB2SCAfFjg0tBKLEUHEc7453B8/urGbGCb5VcKxUqSqteihSbz1UbZFA6oV+mrRULZX4Cv3P7Hq9Y4/BkFSlAh+88/j934+dsS9euhszdESEpDxpebXzVQ+RxOcBTcKWd7PfO7fhIalwEmDGE9LyJkR6l7Y//OAi3lIRiQkC+kRu4ZYXKZVuVSrSh2Usz/OUJLA35CLGCqYirAQCHwPfmFXq1ep6JcY08VCCY2B7YzikPkF9zdLbnjLvMvhKlNQLPhOHmjWxKAw2GNX0Q05khwl0hFnLAzkBP+6Tu8pDDEsFGy2vaj5eZftipSBiagltia5nPjldThCM6oZOhIOCsNZrbF7YLfgbAFOLuG632+nWCn4GgH0fLM10KWMbvY1ae8qzBMqGi7w71Wa1YeNL/NcW8Jvtdru5aeENKBs2FvAb1fXGTt3CG1A2bC7q397pdNYtvAFlw/UFfO/C5nrDxhtQxGgyWkDreBaRKSBDzq444RsA35gmwAxVKWVXRp+oZbkW4ztc9ABggosVTZCapGSIfcB1cDwQFGsBeIvg0k625MuFJS0LSV/QVLW8j1MMFTGDvHj644unj9HJvScn9345uX//5N7PDqorOAnLVM+//+Lvh5+ivx5/9/zBV268LON//+mz33790g1UZeCzrx/98eTRs28+//OHBw74jsCDMrxPYyLRdXKMDngMhjkEkIF4PYp+hGmZYicJJU6wpnGguyqy0NcnmOXRsXBtYnvwloAW4AJeHt+xFD6MxFhRB/BqFFvAPc5ZmwunTVe1rLIXxknoFi7GZdwBxkcu2Z25+HbHKeTyNC1taEQsNfcZhByHJCEK6T0+IsRBdptSy6971Bdc8qFCtylqY+p0SZ8OrGyaEV2hMcRl4lIQ4m35Zu8WanPmYr9LjmwkVAVmLpaEWW68jMcKx06NcczKyGtYRS4lDyfCtxwuFUQ6JIyjbkCkdNHcEBNL3asYepEz7HtsEttIoejIhbyGOS8jd/moE+E4depMk6iM/UiOIEUx2ufKqQS3K0TPIQ44WRruW5RY4X51bd+koaXSLEH0zljkfdvqwDFNXtaOGYV+fNbtGBrgs28f/o8a8Q68k1yVMN9+l+Hmm26Hi4C+/T13F4+TfQJp/r7lvm+572LLXVbPqzbaWW81x+Xpodjwi5eekIeUsUM1YeSaNF1ZgtJBDxbNxBAVB/I0gmEuzsKFApsxElx9QlV0GOEUxNSMhFDmrEOJUi7hGmCWnbz1BrwVVLbWnF4AAY3VHg+y5bXyxbBgY2ahuXxOBa1pBqsKW7twOmG1DLiitJpRbVFaYbJTmnnk3oRqQFhf+2vr9Uw0ZAxmJNB+zxhMw3LmIZIRDkgeI233oiE147cV3KYveatL29RsTyFtlSCVxTWWiJtG7zRRmjKYRUnX7Vw5ssSeoWPQqllvesjHacsbwiEKhnEK/KRuQJiFScvzVW7KK4t53mB3WtaqSw22RKRCql0so4zKbOVELJnpX282tB/OxgBHN1pNi7WN2n+ohXmUQ0uGQ+KrJSuzab7Hx4qIwyg4RgM2FgcY9NapCvYEVMI7w+SangioULMDM7vy8yqY/30mrw7M0gjnPUmX6NTCDG7GhQ5mVlKvmM3p/oammJI/I1PKafyOmaIzF46ta4Ee+nAMEBjpHG15XKiIQxdKI+r3BBwcjCzQC0FZaJUQ0782a13J0axvZTxMQcE5RB3QEAkKnU5FgpB9ldv5Cma1vCvmlZEzyvtMoa5Ms+eAHBHW19W7ru33UDTtJrkjDG4+aPY8d8Yg1IX6tp58srR53ePBTFBGv6qwUtMvvQo2T6fCa75qs461IK7eXPlVm8LlA+kvaNxU+Gx2vu3zA4g+YtMTJYJEPJcdPJAuxWw0AJ2zxUyaZpVJ+LeOUbMQFHLnnF0ujjN0dnFcmnP2y8W9ubPzkeXrch45XF1ZLNFK6SJjZgv/OvHBHZC9CxelMVPS2EfuwlWzM/2/APhkEg3p9j8AAAD//wMAUEsDBBQABgAIAAAAIQD7u/x/EgUAANEPAAARAAAAd29yZC9zZXR0aW5ncy54bWy0V21v2zYQ/j5g/8Hw5znWu2yhTqEXa03RtMOcYp9pibaJSKJAUnHcYf99R0q0nYQp4m79YlP33D13PJ6Op3fvH+tq9IAZJ7RZjO0razzCTUFL0mwX4693+WQ2HnGBmhJVtMGL8QHz8fvrX395t484FgLU+AgoGh7VxWK8E6KNplNe7HCN+BVtcQPghrIaCXhk22mN2H3XTgpat0iQNamIOEwdywrGAw1djDvWRAPFpCYFo5xuhDSJ6GZDCjz8aQv2Fr+9SUaLrsaNUB6nDFcQA234jrRcs9U/ygbgTpM8fG8TD3Wl9fa29Ybt7ikrjxZvCU8atIwWmHM4oLrSAZLm5Nh7QXT0fQW+hy0qKjC3LbU6j9y/jMB5QRAU+PEyjtnAMQXLcx5SXsYTHHnIKbF28GPBnBGU3UUUjqvjkH/S/IyLl6LcXUanz2gqbZFAO8SPFdkzbqrLGL0zxr7AKlrcn3Piy5LmHwkP9ekM+cuwDFXdQ5/ImiHW94yhpOsiutk2lKF1BeFAaY+gOkcqOvkLhyz/1BI/KrnM7bDYVHIBqb+GlvaN0nq0j1rMCnivoR9a1ngqAXib6GYlkADGiLe4qlSDLCqMIIB9tGWohtamJcqmxBvUVeIOrVeCtqD0gGCfoTNQFjvEUCEwW7WoALaUNoLRSuuV9DMVKbRJBm/xYKGa5mm16hswWDSohp0/aaq3tMQyso6Rtx+RNFDebf/c5XNHFC4MRkp8JzO+EocK5xD8inzDcVN+7LggwKha63+I4HsB4EZ6/gI1cndocY6R6CBNP8mZOom8Iu0tYYyym6aE2vhpzshmgxk4IFBrt1A+hNG9yvMHjEq4p3+S347jv0AZXlH3DsryPqFC0PrDod1Brv+Hk5yely9MGyXXiz8pFVrVslLPni3nfaQSPSHW3HGs1IikTuYFRiTzcnvY91PEtrwsH4r8GRK4iZsZkZnnZzMT4oS2ayVGBCJYhibEtew8XRoRxwkzYwRuHLix0Y9nObbrmhHfToxRe7PQs4y59q0g9I0R+KnnJkYkmHneMjYhoeUvQ+MpzOZuEBrPdO5ar7DNY/uVnc7jMLCMGY3jMHYdI5KE8czIlsjyMUaQeLY/G7r4M8S3UnOukzCY+8Zcv17xmefZiXE/2dwLl8a8ZYmfp8bzWULFubkRgbIOjX7y2Mlz5WfaQ/DW1pGcdP9geiVb/6juLVJUrxlBo1s5C0+lxprdJ6TR+BrDhY7PkVW31uBk0gO8RlWVQxPSgEp1HZWEtxneqHV1i9j2xDtoMKMU7uGPRy55r2P2O6Nd26N7htq+pWsV2/MGS9KIT6TWct6tV9qqgRHkDOqa8ssDU3k6pWcfCWiR6mr8hFSrVbq4mXxdyeaIERcxJ2gx/rabpJ/7/BcVW8nOim9R2/YNer21F+OKbHfClmYCnkr4ilIP660zYI7CnB5TD6iQmwXtYXGSOVp2pudqmXuSeVrmnWS+lvknWaBlgZTt4EpmMB/dw12hl1K+oVVF97j8cMJfiPok8B1qcdaPT1BxtBcM8xQfPUT4EYYzXBIBH6ctKWv0KGc1R90Bg3aFDrQTT3QlJpXbpwxyVNa31BNjVfXPYpFjXUGgQleHen2a1q76wCvC4WZtYbATlGnsN4XZXlTS4gZeLlgpue/60DWHtmT7R9jv4b+TJMjmzjKdWLNZOPHiNJ3ES9uZOGGyjD0rSEJ/+c/wbupv8et/AQAA//8DAFBLAwQUAAYACAAAACEAt/GuJK8AAAAOAQAAEwAoAGN1c3RvbVhtbC9pdGVtMS54bWwgoiQAKKAgAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAArI/BCsIwEER/JezdpnoQKW2lIJ5EhCp48JKm2zaQ7JYkiv69QcQv8Dhv4A1Tbp/Oigf6YJgqWGY5CCTNvaGxgst5v9iACFFRrywTVkAM27rsipbvXmMQLVrUEfs2vmyqb82pya7tAcQHHJVLMDEQaYdC0VUwxTgXUgY9oVMh4xkpdQN7p2KKfpQ8DEbjjvXdIUW5yvO17ExnDY9ezdPrK/uLqi7l70z9BgAA//8DAFBLAwQUAAYACAAAACEAL+E+iN8AAABVAQAAGAAoAGN1c3RvbVhtbC9pdGVtUHJvcHMxLnhtbCCiJAAooCAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAACckMFqwzAMQO+D/YPRPXVWmnYrcUo7J9Dr2KBX11ESQ2wF2xkbY/8+h5264y4ST0J6QuXhw47sHX0w5AQ8rHJg6DS1xvUC3l6b7BFYiMq1aiSHAhzBobq/K9uwb1VUIZLHc0TLUsGkfJYCvopit3la18/Zpjim0EiZnepTk8mtbI6yqHeJvoEltUtrgoAhxmnPedADWhVWNKFLzY68VTGh7zl1ndEoSc8WXeTrPN9yPSe9vdgRquWe3+kX7MItLqfN3vzXcjXX0VDv1TR8Aq9K/ke18M0rqh8AAAD//wMAUEsDBBQABgAIAAAAIQD2ObcZYwUAABpFAAASAAAAd29yZC9udW1iZXJpbmcueG1s7JvbbuM2EIbvC/QdDAG9TESdqAPWWThxXGyxLYo2fQBGom0hoiRQtJ28/ZLUwbZkeyV5k2IB3kQ2D59myJnhDzn69PmVJJMtpkWcpVPNuAXaBKdhFsXpaqr997S48bRJwVAaoSRL8VR7w4X2+e7XXz7tgnRDnjHlAyeckRbBLg+n2pqxPND1IlxjgopbEoc0K7Iluw0zomfLZRxifZfRSDeBAeSnnGYhLgrOeUDpFhVahQtf+9EiinZ8sgDaerhGlOHXPcMYDHF0X/e6IHMEiHtoGl2UNRgFdWFVB2SPAnGrOiRnHOmEc3AcyeyS3HEkq0vyxpE64US6AZ7lOOWdy4wSxPhXutIJoi+b/IaDc8Ti5ziJ2RtnAlhjUJy+jLCIz2oIxIoGE1ydZBFOrKimZFNtQ9Ogmn/TzBemB+X86tLMwEm/2/Lb+Tp+ZUnB6rm0z9qV0+dZuCE4ZXLVdIoTvo5ZWqzjvKkOZCyNd65ryPbSAmxJUo/b5UbPVDtX2ublNuyBfcyv9o4kpeWXiQbosZsC0czoY8LxPWtLCI/g/Y1HLc3B4ho9i08NMDsAGOKeh0XN8CqGHu6zW3DinmlVc8pdEZx4v7BGzxrYNuYAEG0GIUyrtkNcxPQDVhGxaD0MV++RLuYihtaoaJKmJC57FoKaaB8QywBLsrCpZ4KJhy2a0wDfyMEe5qvrEvV3mm3yPS2+jvZlX7J3Qj0NYFUJf1iEiuuM+XeNcl7JSRh8WaUZRc8Jt4in74Rn4ETugPjLA1lc5Ef8KttF/FQflon4EG0moiRqd1wFoueCURSyvzZkcvTtC08lriY5PKCYS0gqGkvBOFsyTO8pRi9iiKCkhbhtsEU8rABYQMDTXNNFD9kkLP6Ktzh5estxPWb99kzj6E/Rl4i+ciwjeVKPmJm+7UP4WPYkW9ER80tpVMDyhB/mwAY+4PeTNkgb6+lGOY9r3AVpGiMcxgRVN+OsJ37C1X2/GbdN+x9h3ZrgJSub87+puMSp8FM0TzXXlKasUbqSctuCQIzVq8G6ZLWtN9rWG75s4QccPyW3WIzo502S7TD9ihnfi9MemYM9Mmx7hEtmx6X7a1z6JyMoPe2RdcojGq/W510yDXjskuH1cMk6EWPjXLoYc/bgHTI9b8QO2R8XdM5gl7gHI1xyPizo4PCgs61WaegVdPBjgs4dvEMOGFMW3I8LOm+4S26rLPRyyfuwoPOHBx20W6XhTNDpR8e8oFzUAOLAGq4BfMu2vblXGqs0wLH1SgNckRlKAygNoDSA0gBtl5QGeD8NIKr7YA1gmN7CmDnz0tixGuDB8x5dZ243S9FsxM+iAXYBrS6LLGWFGBmnTFixRNzxaqgco+SCkgtKLnzfJSUXlFy4MuiUXHg/uSBK4WC5YDlw7tgPs9LYsXJhYXsP9v2sohxuhHpk0DMzlAZQGkBpAKUBlAZQGmC8BhC1Y7AGcJwHF/rulRrAX0ATLBawWYpmI9QjAyUXlFz4/g4puaDkgpILbZeUXHg/uSASbbhcmIGZO7u/8hcGY+Z7Mxf+xI8M3kUu/KCk+T/lwY9JEiUHlBxQckDJgbZLSg68nxwQUTlYDriPtgMfHfULwskwUo8ErsgMpQGUBlAaQGmAtktKA1ylAVJ59qcHLxqI9xeDaCPfbpSNjglc4EC/9O9IJtTWyQf9uuR0oPI/F9tQ07CAB6BvnmfWDpxiyn9v6BgKTI6EJrDOQ+0LUPl7SQcKHcdxget656HyPcszUPlUpQ11AXAM1zQ773/umdYFppRm3W2yXGB5oqCehcpwPAMVKduBep7vmtDpvqi6ZzqHzPJaKsW7bwAAAP//AwBQSwMEFAAGAAgAAAAhAIsTHYqQDAAASXgAAA8AAAB3b3JkL3N0eWxlcy54bWy8nVlz2zgSx9+3ar8DS0+7D4l8yHaSGmfKduK1a+PEEzmTZ4iELGxIQktSPubTLwAeAt0ExQZ7/WJLpPqHo/vfAHj+9vtTEgcPPMuFTE8n+2/3JgFPQxmJ9P508uPu8s27SZAXLI1YLFN+Onnm+eT3j3//22+PH/LiOeZ5oABp/iEJTyerolh/mE7zcMUTlr+Va56qnUuZJaxQX7P7acKyX5v1m1Ama1aIhYhF8Tw92Ns7nlSYbAhFLpci5J9kuEl4Whj7acZjRZRpvhLrvKY9DqE9yixaZzLkea4ancQlL2EibTD7MwBKRJjJXC6Lt6oxVY0MSpnv75lPSbwFHOEABwBwHPInHONdxZgqS5sjIhznuOGIyOL4VcYCRBsU4uCwrof+p80tVh4V0QqHq3001basYCuWr9rEZYwjzixiGWCxDH/ZTI7rtKMG+JxoHybhh+v7VGZsESuSispABVZgwPqv8o/+Zz7yJ7Ndd0v1YRnrD6rXPirpRjL8xJdsExe5/prdZtXX6pv5dynTIg8eP7A8FOJO1VcVmghV/tVZmouJ2sNZXpzlgnXuXOkPnXvCvLA2n4tITKa6xPwvtfOBqW4/OKi3XOgatLbFLL2vt/H0zY+5XRNr00JxTycsezM/04bTqmHlf6u565ffTMFrFgpTDlsWXGUllRQ0NBY6CR6cHNdfvm+0L9imkFUhBlD+b7BT0OMqWanUNS8zqNrLl19UrPBoXqgdpxNTltr44/o2EzJTWfJ08v59tXHOE3Elooin1g/TlYj4zxVPf+Q82m7/49IEYrUhlJtUfT5U1TeVyKPPTyFf67yp9qZM++SrNoj1rzdiW7gx/28N26880WW/4kwPHsH+S4SpPgpxoC1yq7XdzM2LtptfoQo6fK2CZq9V0NFrFWSE8BoFnbxWQe9eqyCD+X8WJNJIjQPm97AYQN3FcagRzXGIDc1xaAnNcUgFzXEoAc1xBDqa44hjNMcRpghOIUNXFFrBfuiI9n7u7jHCj7t7SPDj7h4B/Li7E74fd3d+9+PuTud+3N3Z24+7O1njueVUK7hWMkuL0SpbSlmksuBBwZ/G01iqWGZFTcPTgx7PSBpJgCkzWzUQj6aFzHzfHSFGpP7jeaEXfoFcBktxv8l4PrriPH3gsVzzgEWR4hECM15sMkeP+MR0xpc842nIKQObDqpXgkG6SRYEsblm92QsnkbE3VcTSZJCE9Bq/bzSIhEEQZ2wMJPjqyYZWX74IvLxfaUhwfkmjjkR6ytNiBnW+LWBwYxfGhjM+JWBwYxfGFg+o+qiikbUUxWNqMMqGlG/lfFJ1W8VjajfKhpRv1W08f12J4rYpHh71rE//NjdRSz1OZDR9ZiL+5SpCcD44aY6ZhrcsozdZ2y9CvRR6W6s3WZsOecyeg7uKMa0hkQ1rzchcqFaLdLN+A5t0ajE1fCI5NXwiATW8MZL7EZNk/UE7YpmPTPfLIpO0RrSINHOWbwpJ7Tj1caK8RG2FcClyHIyGXRjCSL4q57OandSZL5tLcdXbMsaL6uXWYm0ehWSoJb6hClNGr56XvNMLct+jSZdyjiWjzyiI86LTJaxZkv+wLhkkOQ/J+sVy4VZK7UQw4f6+uqJ4IatRzfoNmYipfHb5zcJE3FAN4O4urv5EtzJtV5m6o6hAZ7LopAJGbM6EviPn3zxT5oKnqlFcPpM1NozosNDBnYhCAaZkiQjIpKaZopUkIyhhvdv/ryQLItoaLcZL68nKTgRcc6SdTnpINCWyouPKv8QzIYM70+WCX1ciEpUdyQw67Bhvln8h4fjU91XGZAcGfq2KczxRzPVNdZ0uPHThBZu/BTBeFMNDzp+CRrbwo1vbAtH1diLmOW5cJ5C9eZRNbfmUbd3/OKv4slYZstNTNeBNZCsB2sgWRfKeJOkOWWLDY+wwYZH3V7CkDE8gkNyhvevTERkzjAwKk8YGJUbDIzKBwZG6oDxV+hYsPGX6Viw8dfqlDCiKYAFo4oz0uGf6CyPBaOKMwOjijMDo4ozA6OKs8NPAV8u1SSYboixkFQxZyHpBpq04MlaZix7JkJ+jvk9IzhAWtJuM7nUd7LItLyImwCpj1HHhJPtEkfl5J98QVY1zaKsF8ERURbHUhIdW9sOOMbSOnB49H6nmbmTY3QVbmMW8pWMI5452uS2VevleXlbxsvqm2oMOuz5RdyvimC+ao7225jjvZ2W9YK9Zba7wK4+P67vZ+kyu+GR2CR1ReHNFMeHw41NRLeMZ7uNtzOJluXRQEtY5vFuy+0suWV5MtASlvluoKXRacuyTw+fWParMxBO+uKnWeM5gu+kL4oa485i+wKpsewKwZO+KGpJJTgLQ322AHpnmGbc9sPE47bHqMhNwcjJTRmsKzeiT2Df+YPQIzsmaZrymqsnXhZ3aCbRgzLnHxtZHrdvnXAaflPXtZo4pTkPOjmHw09ctbKMux8Hpxs3YnDecSMGJyA3YlAmcpqjUpKbMjg3uRGDk5Qbgc5WcETAZStoj8tW0N4nW0GKT7YaMQtwIwZPB9wItFAhAi3UETMFNwIlVGDuJVRIQQsVItBChQi0UOEEDCdUaI8TKrT3ESqk+AgVUtBChQi0UCECLVSIQAsVItBC9ZzbO829hAopaKFCBFqoEIEWqpkvjhAqtMcJFdr7CBVSfIQKKWihQgRaqBCBFipEoIUKEWihQgRKqMDcS6iQghYqRKCFChFooZa3GvoLFdrjhArtfYQKKT5ChRS0UCECLVSIQAsVItBChQi0UCECJVRg7iVUSEELFSLQQoUItFDNycIRQoX2OKFCex+hQoqPUCEFLVSIQAsVItBChQi0UCECLVSIQAkVmHsJFVLQQoUItFAhoi8+q1OUrsvs9/FHPZ1X7A8/dVVV6rt9K7eNOhyOqmvlZg2/F+Fcyl9B542Hh2a9MQwiFrGQ5hC147S6zTWXRKBOfH676L/Dx6aPfOhSdS+EOWcK4LOhluCYyqwv5G1LsMib9UW6bQlmnbO+7GtbgmFw1pd0jS7ri1LUcASM+9KMZbzvMO/L1pY57OK+HG0Zwh7uy8yWIezgvnxsGR4FOjm/tD4a2E/HzfWlgNAXjhbhxE3oC0voqzodQ2EMdZqbMNR7bsJQN7oJKH86MXjHulFoD7tRfq6GMsO62l+obgLW1ZDg5WqA8Xc1RHm7GqL8XA0TI9bVkIB1tX9ydhO8XA0w/q6GKG9XQ5Sfq+FQhnU1JGBdDQlYV48ckJ0Yf1dDlLerIcrP1XByh3U1JGBdDQlYV0OCl6sBxt/VEOXtaojyczVYJaNdDQlYV0MC1tWQ4OVqgPF3NUR5uxqi+lxtjqK0XI3ysGWOm4RZhrgB2TLEJWfL0GO1ZFl7rpYsgudqCfqq9jlutWQ7zU0Y6j03Yagb3QSUP50YvGPdKLSH3Sg/V+NWS12u9heqm4B1NW615HQ1brXU62rcaqnX1bjVktvVuNVSl6txq6UuV/snZzfBy9W41VKvq3GrpV5X41ZLblfjVktdrsatlrpcjVstdbl65IDsxPi7Grda6nU1brXkdjVutdTlatxqqcvVuNVSl6txqyWnq3GrpV5X41ZLva7GrZbcrsatlrpcjVstdbkat1rqcjVuteR0NW611Otq3Gqp19W41dKNMhEEj4CaJywrArrnxV2xfFWw8Q8n/JFmPJfxA48C2qZ+QbVy+th6/ZVmm1cJqt8Xqs/0E9Ct25Wi8gmwFdD88DpqXlOljXVNguqFYNVmU+HqdK35nOVqTV39Zm/vYrb/7nMV5q4Xftmv+5o1X7pf9+V4Z9rp5E4kPA++8sfgu0yY6e7tm8o6dpp3pnXuCXO42dTfemladbq29dI0s81695npQNjl4Ur1eVg9w8vR5dWzeJubycyTeF86wPHAXlOxrRDrX1ehtY2b8netqCnr76h3oYXfU2eTGHpjpcwdrgq+r6JkVw1VfRZxGQPqw3Wqg+2xCqKyptETK1Fq/wWP4xtW/lqu3T+N+bIo9+7vmUc3vNi/KJ9C6LTPzHDlBEzblSm/9sdJ+V6C6joKpzR1Tu7obnNRz9iedtetlTaa2ugyt/c/vqyUGT22u8teZaqkbzrHgZQCq31YiQ6VbISJD+1dfS1TNVCG+tEST8WGxdVd7laCGSCAdsjrY4igseUoaXZ1NdTWiqO19XMC2i3c35t9uqyu+KBJp7aYzmUW8cwME6VYTKn6OeVVw/9ScxbzQZXJmzcdqonAltxIycu2kZmXdS1CL2OhxoOIX40z/9PPvMwHTfd3pYf6U/7xfwAAAP//AwBQSwMEFAAGAAgAAAAhANyn8URlAQAAIAQAABQAAAB3b3JkL3dlYlNldHRpbmdzLnhtbJzTUW/CIBAA4Pcl+w8N70p1akxja7IsJnve9gMQrpYIXAO46n79oFZX44vdSzna3pcDjtX6qFXyDdZJNDmZjFOSgOEopNnl5OtzM1qSxHlmBFNoICcncGRdPD+tmqyB7Qd4H/50SVCMyzTPSeV9nVHqeAWauTHWYMLHEq1mPkztjmpm94d6xFHXzMutVNKf6DRNF6Rj7CMKlqXk8Ib8oMH4Np9aUEFE4ypZu4vWPKI1aEVtkYNzYT1anT3NpLkyk9kdpCW36LD047CYrqKWCumTtI20+gPmw4DpHbDgcBxmLDuDhsy+I8UwZ3F1pOg5/yumB4jDIGL6cqkjDjG9ZznhRTWMu5wRjbnMs4q56lYs1TBx1hPPDaaQ7/smDNu0+RU86XiGmmfvO4OWbVWQQlcmobGSFo7PcD5xaEM4tu/jtnRBqWIQdq0I9xdrL7X8gQ3aV4uNA0uLFb2518UvAAAA//8DAFBLAwQUAAYACAAAACEApdoeknICAABGCAAAEgAAAHdvcmQvZm9udFRhYmxlLnhtbLyUT3KbMBTG953pHRj2MQJj4niCM9M0nukmiyY5gCyE0UR/GEnY8Rm67D16g96mvUefBNh0nDQhi5qxeTzp/dD79MmXV0+CB1uqDVMyD+MJCgMqiSqY3OThw/3qbB4GxmJZYK4kzcM9NeHV8uOHy92iVNKaAOqlWQiSh5W19SKKDKmowGaiaiphsFRaYAuPehMJrB+b+owoUWPL1owzu48ShLKww+i3UFRZMkI/K9IIKq2vjzTlQFTSVKw2PW33FtpO6aLWilBjoGfBW57ATB4wcXoCEoxoZVRpJ9BMtyKPgvIY+UjwI2A2DpCcADJCn8Yx5h0jgsohhxXjONmBw4oB532LGQCKZhQimfbrcDdXPmCZwhbVOFy/R5GrxRZX2FR/E0s+jpgOiK3BuCKPQyYdJ9rsANwLt4eCLL5spNJ4zYEErgzAWIEHu1/YH3fzIX3yeSdLF5TcBaDasju5wW4hsQDQPRPUBLd0F3xVAks/ocZSGRrDnC0GGZATLENTNEMpfBOI0jByE0mFtaEO1k5EbbrEgvF9n9We6wdqZknV57dYM9dMO2TYBgYas0Z5eIMQSm5Wq7DNxHl4DZnz+exTl0ncu/znostMDxnkMsRz/GPccojnHObAO6NWiRNFrjFna81eUGLlFXBXCjoko5QwO2bMOCXSEyXgDypJz/+LEndM3DWtJTC3t5DrV/z7x7dfP793rZxoFINbEKgT99ezGs2z5zTCjVWjJPJdgCePEiXz+apvfyhRnL0iUeomjTWLAK/gF8zijkt7bNzxGWeW9x0blA3NknqzHDJOieTY97/NcvGaWbrALP8AAAD//wMAUEsDBBQABgAIAAAAIQBtMC69gAEAAOoCAAARAAgBZG9jUHJvcHMvY29yZS54bWwgogQBKKAAAQAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAB8kstOwzAQRfdI/EPkfWo7FYVGaSoegg1ISBSB2Bl7WgyJY9lTQv4eJ2lTghC7edw5Gt9xtvwqi+gTnNeVWRA+YSQCIyulzWZBHlfX8RmJPAqjRFEZWJAGPFnmx0eZtKmsHNy7yoJDDT4KJONTaRfkDdGmlHr5BqXwk6AwobmuXCkwpG5DrZAfYgM0YWxGS0ChBAraAmM7EMkOqeSAtFtXdAAlKRRQgkFP+YTTgxbBlf7Pga7zQ1lqbCz8Kd03B/WX14OwrutJPe2kYX9On+9uH7qnxtq0XkkgeaZkihoLyDN6CEPkt6/vILEvD0mIpQOBlcuv4NWJrruvtF5/QFNXTvkwN8qCTIGXTlsMF+ypo0JQF8LjXTjpWoO6aPLzcAZhopttIbDpcL8E7YyDT93+iZzPOsmQZzuH++1ARcGZtPdx33maXl6trkmeMD6P2Tzmpyt2mibTlLGXdsHR/AFY7jb4l5jMYpbEnK0YT6cnY+Ie0Hs0/p35NwAAAP//AwBQSwMEFAAGAAgAAAAhAGrLGZHjAQAA5AMAABAACAFkb2NQcm9wcy9hcHAueG1sIKIEASigAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAnFPLbtswELwX6D8IuseU7NQNDJpB4aDIoW0MWEnOLLWyiVIkQW6MuF/fpRSrdNtTddoZroazD/Lb194URwhRO7su61lVFmCVa7Xdr8vH5vPVTVlElLaVxllYlyeI5a14/45vg/MQUEMsSMLGdXlA9CvGojpAL+OMji2ddC70EgmGPXNdpxXcOfXSg0U2r6olg1cE20J75SfBclRcHfF/RVunkr/41Jw86QneQO+NRBDf0p9m1jrsOZtY3jiUptE9iMWC+AnxrdxDFHPOxoA/u9ASXi45G0O+OcggFVILRf1hXnOWEfyT90YridRd8VWr4KLrsHgYLBdJgLM8hVMZO1AvQeNJVJzlkH/RlhzU5GWMyFuQ+yD9IYpFMjghvlPSwIY6IDppInD2m+D3INN0t1Ing0dcHUGhC0XUP2m+87L4LiOkvq3LowxaWizHtBEMsfERg2g0GtKe8BDmaXmsr0U9JFBwmTiAwQPFl+6GG+JDR7XhP8zWudnBw2g1s5M7O9/xh+rG9V7aE62GpWrOKLX7R3z0jbtLG/LWx0sym/2zxsPOS5Um9PHmOt+C7IjviIWWxjoNZiL4PZURTLqA/rV7aM85fx+kvXoaH62ol7OKvmGRzhxtw/SaxC8AAAD//wMAUEsDBBQABgAIAAAAIQB0Pzl6wgAAACgBAAAeAAgBY3VzdG9tWG1sL19yZWxzL2l0ZW0xLnhtbC5yZWxzIKIEASigAAEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAjM+xisMwDAbg/eDewWhvnNxQyhGnSyl0O0oOuhpHSUxjy1hqad++5qYrdOgoif/7Ubu9hUVdMbOnaKCpalAYHQ0+TgZ++/1qA4rFxsEuFNHAHRm23edHe8TFSgnx7BOrokQ2MIukb63ZzRgsV5QwlstIOVgpY550su5sJ9Rfdb3W+b8B3ZOpDoOBfBgaUP094Ts2jaN3uCN3CRjlRYV2FxYKp7D8ZCqNqrd5QjHgBcPfqqmKCbpr9dN/3QMAAP//AwBQSwECLQAUAAYACAAAACEA5CSJTH0BAAApBgAAEwAAAAAAAAAAAAAAAAAAAAAAW0NvbnRlbnRfVHlwZXNdLnhtbFBLAQItABQABgAIAAAAIQAekRq37wAAAE4CAAALAAAAAAAAAAAAAAAAALYDAABfcmVscy8ucmVsc1BLAQItABQABgAIAAAAIQBQy797UAsAAKo9AAARAAAAAAAAAAAAAAAAANYGAAB3b3JkL2RvY3VtZW50LnhtbFBLAQItABQABgAIAAAAIQDX5gieGwEAAEQEAAAcAAAAAAAAAAAAAAAAAFUSAAB3b3JkL19yZWxzL2RvY3VtZW50LnhtbC5yZWxzUEsBAi0AFAAGAAgAAAAhAJa1reLxBQAAUBsAABUAAAAAAAAAAAAAAAAAshQAAHdvcmQvdGhlbWUvdGhlbWUxLnhtbFBLAQItABQABgAIAAAAIQD7u/x/EgUAANEPAAARAAAAAAAAAAAAAAAAANYaAAB3b3JkL3NldHRpbmdzLnhtbFBLAQItABQABgAIAAAAIQC38a4krwAAAA4BAAATAAAAAAAAAAAAAAAAABcgAABjdXN0b21YbWwvaXRlbTEueG1sUEsBAi0AFAAGAAgAAAAhAC/hPojfAAAAVQEAABgAAAAAAAAAAAAAAAAAHyEAAGN1c3RvbVhtbC9pdGVtUHJvcHMxLnhtbFBLAQItABQABgAIAAAAIQD2ObcZYwUAABpFAAASAAAAAAAAAAAAAAAAAFwiAAB3b3JkL251bWJlcmluZy54bWxQSwECLQAUAAYACAAAACEAixMdipAMAABJeAAADwAAAAAAAAAAAAAAAADvJwAAd29yZC9zdHlsZXMueG1sUEsBAi0AFAAGAAgAAAAhANyn8URlAQAAIAQAABQAAAAAAAAAAAAAAAAArDQAAHdvcmQvd2ViU2V0dGluZ3MueG1sUEsBAi0AFAAGAAgAAAAhAKXaHpJyAgAARggAABIAAAAAAAAAAAAAAAAAQzYAAHdvcmQvZm9udFRhYmxlLnhtbFBLAQItABQABgAIAAAAIQBtMC69gAEAAOoCAAARAAAAAAAAAAAAAAAAAOU4AABkb2NQcm9wcy9jb3JlLnhtbFBLAQItABQABgAIAAAAIQBqyxmR4wEAAOQDAAAQAAAAAAAAAAAAAAAAAJw7AABkb2NQcm9wcy9hcHAueG1sUEsBAi0AFAAGAAgAAAAhAHQ/OXrCAAAAKAEAAB4AAAAAAAAAAAAAAAAAtT4AAGN1c3RvbVhtbC9fcmVscy9pdGVtMS54bWwucmVsc1BLBQYAAAAADwAPANQDAAC7QAAAAAA='''


# I/O
def load_json_files(glob_pattern: str):
    files = sorted(glob(glob_pattern))
    data = []
    for f in files:
        try:
            with open(f, "r", encoding="utf-8") as fh:
                js = json.load(fh)
            data.append((f, js))
        except Exception as e:
            print(f"Failed to read {f}: {e}")
    return data


def ensure_template_docx(path: str) -> None:
    """
    Writes the embedded template DOCX to disk if it doesn't exist.
    """
    if os.path.exists(path) and os.path.getsize(path) > 0:
        return

    b64_clean = "".join((TEMPLATE_DOCX_B64 or "").split())
    if not b64_clean:
        raise RuntimeError(
            "TEMPLATE_DOCX_B64 is empty. Paste your template DOCX Base64 string into TEMPLATE_DOCX_B64."
        )

    try:
        raw = base64.b64decode(b64_clean)
    except Exception as e:
        raise RuntimeError(f"Failed to decode TEMPLATE_DOCX_B64: {e}")

    with open(path, "wb") as f:
        f.write(raw)


#JSON NORMALIZATION
def strip_html(s: str) -> str:
    if not s:
        return ""
    s = html.unescape(s)
    s = re.sub(r"<br\s*/?>", "\n", s, flags=re.I)
    s = re.sub(r"<[^>]+>", "", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def get_outcome_title(js: Dict[str, Any]) -> str:
    if js.get("outcome_title"):
        return str(js["outcome_title"]).strip()
    oi = js.get("outcome_identification") or {}
    return str(oi.get("title") or "N/A").strip()


def get_outcome_long_desc(js: Dict[str, Any]) -> str:
    if js.get("outcome_long_description"):
        return strip_html(str(js["outcome_long_description"]))
    oi = js.get("outcome_identification") or {}
    return strip_html(str(oi.get("long_description") or oi.get("description") or ""))


def get_course_name(js: Dict[str, Any]) -> str:
    ci = js.get("course_info") or js.get("course_identification") or {}
    return str(ci.get("name") or "N/A").strip()


def get_course_code(js: Dict[str, Any]) -> str:
    ci = js.get("course_info") or js.get("course_identification") or {}
    return str(ci.get("course_code") or "N/A").strip()


def get_enrollment_count(js: Dict[str, Any]) -> int:
    ci = js.get("course_info") or js.get("course_identification") or {}
    enrollments = ci.get("enrollments", [])
    if isinstance(enrollments, list):
        for enrollment in enrollments:
            if isinstance(enrollment, dict) and enrollment.get("type") == "student":
                amount = enrollment.get("amount")
                if amount is not None:
                    try:
                        return int(amount)
                    except (ValueError, TypeError):
                        pass
    return 0


def normalize_course_line_for_template(course_code: str) -> str:
    """
    Template wants: 'CSE # 301'
    If we see 'CSE 301' or 'CSE301' or 'CSE #301', normalize.
    """
    s = (course_code or "").strip()
    m = re.search(r"\b([A-Z]{2,4})\s*#?\s*(\d{3})\b", s)
    if m:
        return f"{m.group(1)} # {m.group(2)}"
    return s if s else "N/A"


def get_overall_summary(js: Dict[str, Any]) -> Dict[str, Any]:
    """
    Returns normalized overall summary dict with keys:
      sample_size, number_competent, percent_competent, outcome_met
    """
    if isinstance(js.get("assessment_summary"), dict):
        a = js["assessment_summary"]
        return {
            "sample_size": a.get("total_students_assessed"),
            "number_competent": a.get("number_competent"),
            "percent_competent": a.get("percent_competent"),
            "outcome_met": a.get("outcome_met"),
        }

    r = js.get("results") or {}
    o = r.get("overall_summary")
    if isinstance(o, dict):
        return {
            "sample_size": o.get("sample_size"),
            "number_competent": o.get("number_competent"),
            "percent_competent": o.get("percent_competent"),
            "outcome_met": o.get("outcome_met"),
        }

    return {}


def get_threshold(js: Dict[str, Any]) -> str:
    assigns = js.get("contributing_assignments") or []
    first = assigns[0] if isinstance(assigns, list) and assigns and isinstance(assigns[0], dict) else {}
    thr = first.get("threshold")
    if thr is None or thr == "":
        return "70%"
    return str(thr).strip()


def infer_metric_instrument_type(js: Dict[str, Any]) -> str:
    assigns = js.get("contributing_assignments") or []
    first = assigns[0] if isinstance(assigns, list) and assigns and isinstance(assigns[0], dict) else {}
    explicit = first.get("metric_instrument_type") or first.get("metric")
    if explicit:
        return str(explicit).strip()

    name = str(first.get("name") or "").lower()
    desc = str(first.get("description") or "").lower()

    if "essay" in name or "essay" in desc:
        return "Essay"
    if "exam" in name or "exam" in desc:
        return "Exam"
    if "quiz" in name or "quiz" in desc:
        return "Quiz"
    if "project" in name:
        return "Project"
    return "Assignment"


def build_structured_summary(js: Dict[str, Any]) -> str:
    """
    Only used to give the LLM context when outcome was not met.
    """
    overall = get_overall_summary(js)

    assigns = js.get("contributing_assignments") or []
    rubric_count = 0
    if isinstance(assigns, list):
        for a in assigns:
            if isinstance(a, dict) and isinstance(a.get("rubric"), list):
                rubric_count += len(a["rubric"])

    lines = [
        f"Course name: {get_course_name(js)}",
        f"Course code: {get_course_code(js)}",
        f"Outcome title: {get_outcome_title(js)}",
        f"Threshold: {get_threshold(js)}",
        f"Sample size: {overall.get('sample_size', 'N/A')}",
        f"Number competent: {overall.get('number_competent', 'N/A')}",
        f"Percent competent: {overall.get('percent_competent', 'N/A')}",
        f"Outcome met: {overall.get('outcome_met', 'N/A')}",
        f"Rubric criteria total: {rubric_count}",
    ]
    return "\n".join(lines)


# OPENAI (FEEDBACK ONLY)
def generate_feedback_with_openai(summary_text: str) -> str:
    api_key = os.getenv("OPENAI_API_KEY")
    # sanitize common .env formatting issues (surrounding quotes, stray spaces)
    if api_key:
        api_key = api_key.strip().strip('"').strip("'")

    if not api_key:
        print("OPENAI_API_KEY not set; skipping OpenAI feedback and returning 'NA'.")
        return "NA"

    try:
        client = OpenAI(api_key=api_key)
        response = client.responses.create(
            model=OPENAI_MODEL,
            input=[
                {"role": "system", "content": prompt_base},
                {"role": "user", "content": summary_text},
            ],
        )
        return (response.output_text or "").strip() or "NA"
    except Exception as e:
        print(f"OpenAI call failed: {e}")
        return "NA"


# DOCX EDITING (PRESERVE TEMPLATE FORMATTING)
def replace_paragraph_text_preserve_style(p, new_text: str) -> None:
    """
    Replaces paragraph text while preserving formatting of the first run (best effort).
    """
    if p.runs:
        proto = p.runs[0]
        font = proto.font
        bold = font.bold
        italic = font.italic
        underline = font.underline
        name = font.name
        size = font.size
    else:
        bold = italic = underline = name = size = None

    for r in p.runs:
        r.text = ""

    run = p.add_run(new_text)
    if name is not None:
        run.font.name = name
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if underline is not None:
        run.font.underline = underline


def iter_all_paragraphs(doc: Document):
    """
    Yield paragraphs in document body AND inside all table cells.
    """
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def safe_int(x) -> Optional[int]:
    try:
        if x is None:
            return None
        if isinstance(x, bool):
            return None
        if isinstance(x, (int, float)):
            return int(x)
        s = str(x).strip()
        if s == "":
            return None
        return int(float(s))
    except Exception:
        return None


def update_score_distribution_table(doc: Document, thr: str, sample_size, number_competent) -> None:
    """
    Expected template labels (from your template):
      - "100%-70%"  
      - "Below 70%"
    """
    n_total = safe_int(sample_size)
    n_comp = safe_int(number_competent)
    if n_total is None or n_comp is None:
        return
    n_below = max(0, n_total - n_comp)

    # normalize threshold like "70%" -> "70"
    thr_num = re.sub(r"[^0-9.]", "", str(thr))
    thr_label = f"Below {thr_num}%" if thr_num else "Below 70%"
    comp_label = f"100%-{thr_num}%" if thr_num else "100%-70%"

    label_variants = {
        "competent": {comp_label, "100%-70%", "100% - 70%", "70%-100%", "70% - 100%"},
        "below": {thr_label, "Below 70%", "Below70%", "Below 70 %"},
    }

    for table in doc.tables:
        for row in table.rows:
            # collect cell texts
            cells = row.cells
            texts = [(" ".join((p.text or "").strip() for p in c.paragraphs)).strip() for c in cells]

            # find which row this is
            row_type = None
            for t in texts:
                if t in label_variants["competent"]:
                    row_type = "competent"
                    break
                if t in label_variants["below"]:
                    row_type = "below"
                    break
            if not row_type:
                continue

            # Strategy: write into the first purely-numeric-looking cell OR the cell immediately right of the label.
            # 1) find label cell index
            label_idx = None
            for i, t in enumerate(texts):
                if row_type == "competent" and t in label_variants["competent"]:
                    label_idx = i
                    break
                if row_type == "below" and t in label_variants["below"]:
                    label_idx = i
                    break

            new_val = str(n_comp if row_type == "competent" else n_below)

            # 2) try numeric cell
            wrote = False
            for i, t in enumerate(texts):
                if re.fullmatch(r"\d+", t):
                    # replace all paragraphs in that cell with a single preserved-style paragraph
                    cell = cells[i]
                    if cell.paragraphs:
                        replace_paragraph_text_preserve_style(cell.paragraphs[0], new_val)
                        for extra in cell.paragraphs[1:]:
                            replace_paragraph_text_preserve_style(extra, "")
                    else:
                        cell.add_paragraph(new_val)
                    wrote = True
                    break

            # 3) fallback: cell right of label
            if not wrote and label_idx is not None and label_idx + 1 < len(cells):
                cell = cells[label_idx + 1]
                if cell.paragraphs:
                    replace_paragraph_text_preserve_style(cell.paragraphs[0], new_val)
                    for extra in cell.paragraphs[1:]:
                        replace_paragraph_text_preserve_style(extra, "")
                else:
                    cell.add_paragraph(new_val)


def update_section1_in_doc(doc: Document, js: Dict[str, Any], feedback_text: Optional[str]) -> None:
    outcome_title = get_outcome_title(js)
    outcome_long = get_outcome_long_desc(js)
    if outcome_long:
        outcome_line = f"Outcome assessed ({outcome_title}) {outcome_long}"
    else:
        outcome_line = f"Outcome assessed ({outcome_title})"

    class_line = normalize_course_line_for_template(get_course_code(js))
    metric_line = f"Metric Instrument Type: {infer_metric_instrument_type(js)}"
    thr = get_threshold(js)

    enrollment_count = get_enrollment_count(js)
    enrollment_line = f"Class Enrollment: {enrollment_count} students"

    overall = get_overall_summary(js)
    sample_size = overall.get("sample_size", "N/A")
    number_comp = overall.get("number_competent", "N/A")
    pct_comp = overall.get("percent_competent", "N/A")
    outcome_met = overall.get("outcome_met", None)

    sample_line = f"Sample size: {sample_size} students"

    if isinstance(number_comp, (int, float)) and (pct_comp is not None):
        comp_summary = f"{int(number_comp)} students scored {thr} or above on rubric ({pct_comp}%)"
    else:
        comp_summary = f"N/A students scored {thr} or above on rubric (N/A%)"

    if outcome_met is True:
        outcome_status_line = "Outcome was met"
        changes_tail = "NA"
    elif outcome_met is False:
        outcome_status_line = "Outcome was not met"
        changes_tail = (feedback_text or "NA").strip() or "NA"
    else:
        outcome_status_line = "Outcome was N/A"
        changes_tail = "NA"

    changes_prefix = (
        "If outcome was not met, what changes need to be made to ensure that students can meet this outcome in the future?"
    )
    changes_line_full = f"{changes_prefix} {changes_tail}"

    # Replace by matching existing template lines across body + tables
    for p in iter_all_paragraphs(doc):
        t = (p.text or "").strip()

        if t.startswith("Outcome assessed"):
            replace_paragraph_text_preserve_style(p, outcome_line)
            continue

        if re.fullmatch(r"[A-Z]{2,4}\s*#\s*\d{3}", t):
            replace_paragraph_text_preserve_style(p, class_line)
            continue

        if t.startswith("Metric Instrument Type:"):
            replace_paragraph_text_preserve_style(p, metric_line)
            continue

        if t.startswith("Class Enrollment:"):
            replace_paragraph_text_preserve_style(p, enrollment_line)
            continue

        if t.startswith("To show competency, a student must score at least"):
            replace_paragraph_text_preserve_style(
                p, f"To show competency, a student must score at least {thr} on the assessment."
            )
            continue

        if t.startswith("To show the outcome has been met,"):
            replace_paragraph_text_preserve_style(
                p, f"To show the outcome has been met, {thr} of the students must show competency."
            )
            continue

        if t.startswith("Sample size:"):
            replace_paragraph_text_preserve_style(p, sample_line)
            continue

        if "students scored" in t and "or above on rubric" in t:
            replace_paragraph_text_preserve_style(p, comp_summary)
            continue

        if t.startswith("Outcome was"):
            replace_paragraph_text_preserve_style(p, outcome_status_line)
            continue

        if t.startswith(changes_prefix):
            replace_paragraph_text_preserve_style(p, changes_line_full)
            continue

    update_score_distribution_table(doc, thr, sample_size, number_comp)


# main
def main():
    ensure_template_docx(template_path)

    data = load_json_files(json_input_glob)
    summary_records = []

    for path, js in data:
        base = Path(path).stem
        print(f"\nProcessing {base} ...")

        overall = get_overall_summary(js)
        outcome_met = overall.get("outcome_met", None)

        feedback_text = None
        if outcome_met is False:
            summary_text = build_structured_summary(js)
            feedback_text = generate_feedback_with_openai(summary_text)

        doc = Document(template_path)
        update_section1_in_doc(doc, js, feedback_text)

        out_path = os.path.join(out_dir, f"{base}_ABET_Report.docx")
        doc.save(out_path)
        print(f"Word report saved → {out_path}")

        ca = js.get("contributing_assignments", []) or []
        summary_records.append(
            {
                "input_file": path,
                "contributing_assignments": len(ca) if isinstance(ca, list) else 0,
                "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
            }
        )

    if summary_records:
        with open(output_csv, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=summary_records[0].keys())
            writer.writeheader()
            writer.writerows(summary_records)
        print(f"\nSummary CSV written: {output_csv}")


if __name__ == "__main__":
    main()
